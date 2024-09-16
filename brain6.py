import re
import languagemodels as lm
import spacy
from textblob import TextBlob
from db_operations import find_section_info
from docx import Document
from translate import Translator

nlp = spacy.load('en_core_web_sm')

def message_probability(user_message, recognised_words, single_response=False, required_words=[]):
    message_certainty = 0
    has_required_words = True

    for word in user_message:
        if word in recognised_words:
            message_certainty += 1

    percentage = float(message_certainty) / float(len(recognised_words))

    for word in required_words:
        if word not in user_message:
            has_required_words = False
            break

    if has_required_words or single_response:
        return int(percentage * 100)
    else:
        return 0

def check_all_messages(message):
    highest_prob_list = {}

    def response(bot_response, list_of_words, single_response=False, required_words=[]):
        nonlocal highest_prob_list
        highest_prob_list[bot_response] = message_probability(message, list_of_words, single_response, required_words)

    response('Hello!', ['hello', 'hi', 'hey', 'sup', 'heyo'], single_response=True)
    response('See you!', ['bye', 'goodbye'], single_response=True)
    response('I\'m doing fine, and you?', ['how', 'are', 'you'], required_words=['how'])
    response('You\'re welcome!', ['thank', 'thanks'], single_response=True)

    best_match = max(highest_prob_list, key=highest_prob_list.get)

    return None if highest_prob_list[best_match] < 1 else best_match

def extract_entities(input_text):
    doc = nlp(input_text)
    entities = {ent.label_: ent.text for ent in doc.ents}
    return entities

def extract_section_number(entities):
    if 'LAW' in entities:
        input2=entities['LAW']
        words = input2.split()
        for word in words:
            if word.isdigit():  
                return int(word)
    elif 'CARDINAL' in entities:
        input2=entities['CARDINAL']
        words = input2.split()
        for word in words:
            if word.isdigit():  
                return int(word)
    else:
        return None

def analyze_sentiment(text):
    analysis = TextBlob(text)
    return analysis.sentiment.polarity

def generate_bail_application(details):
    doc = Document()
    doc.add_heading('Bail Application', 0)
    
    doc.add_paragraph(f"Applicant: {details['applicant_name']}")
    doc.add_paragraph(f"Section Number: {details['section_number']}")
    doc.add_paragraph(f"Offense Description: {details['offense_description']}")
    doc.add_paragraph(f"Reason for Bail: {details['reason_for_bail']}")
    
    doc.add_paragraph(f"Duration of Imprisonment Already Served: {details['imprisonment_duration']}")
    doc.add_paragraph(f"Risk Assessment: {details['risk_assessment']}")
    
    doc.save('bail_application.docx')
    return 'Bail application document generated successfully.'

def ReplyBrain(input2, user_language='en'):
    split_message = re.split(r'\s+|[,;?!.-]\s*', input2.lower())
    predefined_response = check_all_messages(split_message)


    lm.config["max_ram"] = "4gb"
    
    translator = Translator(to_lang='en')
    translated_input = translator.translate(input2)
    
    sentiment = analyze_sentiment(translated_input)
    
    entities = extract_entities(translated_input)
    
    section_number = extract_section_number(entities)

    if "application" in translated_input.lower():
        applicant_name = input("Enter the applicant name:")
        section_number2 = input("Enter the section number:")
        offense_description = input("Enter the offense description:")
        reason_for_bail = input("Enter the reason for bail:")
        imprisonment_duration = input("Enter the imprisonment duration:")
        risk_assessment = input("Enter the risk assessment:")
        details = {
            'applicant_name': applicant_name,  
            'section_number': section_number2,
            'offense_description': offense_description,
            'reason_for_bail': reason_for_bail,
            'imprisonment_duration': imprisonment_duration,
            'risk_assessment': risk_assessment
        }
        response = generate_bail_application(details)
        translator = Translator(to_lang=user_language)
        translated_response = translator.translate(response)
        return translated_response
    elif predefined_response:
        return predefined_response
    elif section_number:
        section_info = find_section_info(section_number)
        
        if sentiment < 0:
            response = f"I understand this is concerning. {section_number}: {section_info}"
        else:
            response = f" {section_number}: {section_info}"
        
        translator = Translator(to_lang=user_language)
        translated_response = translator.translate(response)
        return translated_response
    else:
        output = lm.do(translated_input)
        
        if sentiment < 0:
            response = "It seems you're feeling worried. Let me help you with that: " + output
        else:
            response = output
    
        translator = Translator(to_lang=user_language)
        translated_response = translator.translate(response)
        return translated_response
